import os
import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document

# PDF dönüşümü için Windows Word COM
import win32com.client


# =========================
# AYARLAR
# =========================
BASE_DIR = Path(__file__).resolve().parent
EXCEL_PATH = BASE_DIR / "liste.xlsx"
TEMPLATE_PATH = BASE_DIR / "soru.docx"
OUTPUT_DIR = BASE_DIR / "cikti"

# Excel sütun adları
NAME_COLUMN_CANDIDATES = ["isim", "adı soyadı", "ad soyad", "adi soyadi", "ad_soyad", "isim soyisim"]
JOB_COLUMN_CANDIDATES = ["meslek", "görev", "gorev", "unvan", "pozisyon"]


# =========================
# YARDIMCI FONKSİYONLAR
# =========================
def normalize_text(value: str) -> str:
    """Sütun eşleştirme için küçük harf ve sadeleştirme yapar."""
    if value is None:
        return ""
    value = str(value).strip().lower()
    replacements = {
        "ı": "i",
        "ğ": "g",
        "ü": "u",
        "ş": "s",
        "ö": "o",
        "ç": "c",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def find_column(df: pd.DataFrame, candidates: list[str]) -> str:
    """Excel'de uygun sütun adını bulur."""
    normalized_map = {normalize_text(col): col for col in df.columns}

    for candidate in candidates:
        norm_candidate = normalize_text(candidate)
        if norm_candidate in normalized_map:
            return normalized_map[norm_candidate]

    raise ValueError(
        f"Uygun sütun bulunamadı. Aranan alternatifler: {candidates}\n"
        f"Excel sütunları: {list(df.columns)}"
    )


def safe_filename(text: str) -> str:
    """Dosya adı için güvenli metin üretir."""
    text = str(text).strip()
    text = re.sub(r'[\\/*?:"<>|]', "_", text)
    text = re.sub(r"\s+", " ", text)
    return text


def replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """
    Paragraf içindeki metni değiştirir.
    Not: Yer tutucular tek parça metin olarak yazılmışsa güvenilir çalışır.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for old, new in replacements.items():
        new_text = new_text.replace(old, new)

    if new_text != full_text:
        # Mevcut run'ları temizle
        for run in paragraph.runs:
            run.text = ""

        # İlk run'a yeni metni yaz
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def replace_text_in_doc(doc: Document, replacements: dict[str, str]) -> None:
    """Belgedeki paragraflar ve tablolar dahil tüm alanlarda değiştirme yapar."""
    # Normal paragraflar
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Tablolar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)


def create_docx_from_template(template_path: Path, output_docx_path: Path, replacements: dict[str, str]) -> None:
    """Şablon docx dosyasını doldurup yeni docx oluşturur."""
    doc = Document(str(template_path))
    replace_text_in_doc(doc, replacements)
    doc.save(str(output_docx_path))


def convert_docx_to_pdf(word_app, input_docx: Path, output_pdf: Path) -> None:
    """
    Word COM ile DOCX -> PDF dönüştürür.
    FileFormat=17 => PDF
    """
    doc = None
    try:
        doc = word_app.Documents.Open(str(input_docx.resolve()))
        doc.SaveAs(str(output_pdf.resolve()), FileFormat=17)
    finally:
        if doc is not None:
            doc.Close(False)


def main():
    # Dosya kontrolleri
    if not EXCEL_PATH.exists():
        raise FileNotFoundError(f"Excel dosyası bulunamadı: {EXCEL_PATH}")

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Şablon dosyası bulunamadı: {TEMPLATE_PATH}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Excel oku
    df = pd.read_excel(EXCEL_PATH)

    if df.empty:
        raise ValueError("Excel dosyası boş görünüyor.")

    # Sütunları bul
    name_col = find_column(df, NAME_COLUMN_CANDIDATES)
    job_col = find_column(df, JOB_COLUMN_CANDIDATES)

    # Word uygulaması başlat
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    created_count = 0

    try:
        for index, row in df.iterrows():
            isim = "" if pd.isna(row[name_col]) else str(row[name_col]).strip()
            meslek = "" if pd.isna(row[job_col]) else str(row[job_col]).strip()

            # Boş satır atla
            if not isim and not meslek:
                continue

            if not isim:
                print(f"Satır {index + 2}: isim boş, atlandı.")
                continue

            replacements = {
                "<isim>": isim,
                "<meslek>": meslek,
            }

            file_base = safe_filename(isim)
            output_docx = OUTPUT_DIR / f"{file_base}.docx"
            output_pdf = OUTPUT_DIR / f"{file_base}.pdf"

            create_docx_from_template(TEMPLATE_PATH, output_docx, replacements)
            convert_docx_to_pdf(word, output_docx, output_pdf)

            created_count += 1
            print(f"Oluşturuldu: {output_docx.name} | {output_pdf.name}")

    finally:
        word.Quit()

    print(f"\nTamamlandı. Toplam {created_count} adet belge üretildi.")
    print(f"Çıktı klasörü: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()